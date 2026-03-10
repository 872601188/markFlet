import flet as ft
import markdown
import sqlite3
import os
import io
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import ImageGrab, Image
from bs4 import BeautifulSoup

# 可选依赖：PDF 导出
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False


class Database:
    """SQLite 数据库管理"""
    
    def __init__(self, db_path="markflet.db"):
        self.db_path = db_path
        self.init_db()
    
    def init_db(self):
        """初始化数据库"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS recent_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_path TEXT UNIQUE,
                file_name TEXT,
                opened_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()
        conn.close()
    
    def add_recent_file(self, file_path):
        """添加最近打开的文件"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        file_name = os.path.basename(file_path)
        cursor.execute('''
            INSERT OR REPLACE INTO recent_files (file_path, file_name, opened_at)
            VALUES (?, ?, CURRENT_TIMESTAMP)
        ''', (file_path, file_name))
        conn.commit()
        conn.close()
    
    def get_recent_files(self, limit=10):
        """获取最近打开的文件列表"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT file_path, file_name, opened_at 
            FROM recent_files 
            ORDER BY opened_at DESC 
            LIMIT ?
        ''', (limit,))
        files = cursor.fetchall()
        conn.close()
        return files


class MarkdownConverter:
    """Markdown 转换器"""
    
    @staticmethod
    def md_to_html(md_text):
        """将 Markdown 转换为 HTML"""
        return markdown.markdown(
            md_text,
            extensions=[
                'pymdownx.extra',
                'pymdownx.highlight',
                'pymdownx.superfences',
                'tables',
                'fenced_code',
            ]
        )
    
    @staticmethod
    def _parse_inline_formatting(paragraph, text):
        """解析行内格式（粗体、斜体、行内代码、链接）"""
        # 处理链接 [text](url)
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
        parts = re.split(link_pattern, text)
        
        i = 0
        while i < len(parts):
            if i + 2 < len(parts) and parts[i+1] is not None:
                # 这是链接部分
                link_text = parts[i]
                link_url = parts[i+1]
                run = paragraph.add_run(link_text)
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
                i += 3
            else:
                # 普通文本，处理格式
                MarkdownConverter._parse_formatting(paragraph, parts[i])
                i += 1
    
    @staticmethod
    def _parse_formatting(paragraph, text):
        """解析粗体、斜体、行内代码"""
        if not text:
            return
            
        # 处理粗体 **text**
        bold_parts = text.split('**')
        for idx, part in enumerate(bold_parts):
            if idx % 2 == 1:  # 粗体部分
                run = paragraph.add_run(part)
                run.bold = True
            else:
                # 处理斜体 *text*
                italic_parts = part.split('*')
                for sub_idx, sub_part in enumerate(italic_parts):
                    if sub_idx % 2 == 1:  # 斜体部分
                        run = paragraph.add_run(sub_part)
                        run.italic = True
                    else:
                        # 处理行内代码 `code`
                        code_parts = sub_part.split('`')
                        for code_idx, code_part in enumerate(code_parts):
                            if code_idx % 2 == 1:  # 代码部分
                                run = paragraph.add_run(code_part)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(128, 128, 128)
                            else:
                                if code_part:
                                    paragraph.add_run(code_part)
    
    @staticmethod
    def md_to_docx(md_text, output_path, base_dir=None):
        """
        将 Markdown 转换为 Word 文档（使用 HTML 中间表示实现更精确的转换）
        """
        doc = Document()
        
        # 将 Markdown 转换为 HTML
        html_content = markdown.markdown(
            md_text,
            extensions=[
                'pymdownx.extra',
                'pymdownx.highlight',
                'pymdownx.superfences',
                'tables',
                'fenced_code',
                'toc'
            ]
        )
        
        # 使用 BeautifulSoup 解析 HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 遍历所有顶级元素
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'ul', 'ol', 'blockquote', 'table', 'img']):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'h4':
                doc.add_heading(element.get_text(), level=4)
            elif element.name == 'h5':
                doc.add_heading(element.get_text(), level=5)
            elif element.name == 'h6':
                doc.add_heading(element.get_text(), level=6)
                
            elif element.name == 'p':
                p = doc.add_paragraph()
                MarkdownConverter._parse_inline_elements(p, element, base_dir)
                
            elif element.name == 'pre':
                # 代码块
                code = element.find('code')
                if code:
                    # 获取语言类型（如果有）
                    classes = code.get('class', [])
                    lang = ''
                    for cls in classes:
                        if cls.startswith('language-'):
                            lang = cls.replace('language-', '')
                            break
                    
                    code_text = code.get_text()
                    p = doc.add_paragraph()
                    if lang:
                        p.add_run(f'📝 {lang}\n').italic = True
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.shading.background_pattern = True
                    
            elif element.name == 'ul':
                # 无序列表
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Bullet')
                    MarkdownConverter._parse_inline_elements(p, li, base_dir)
                    
            elif element.name == 'ol':
                # 有序列表
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Number')
                    MarkdownConverter._parse_inline_elements(p, li, base_dir)
                    
            elif element.name == 'blockquote':
                # 引用块
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                
                # 处理嵌套元素
                for child in element.children:
                    if child.name == 'p':
                        MarkdownConverter._parse_inline_elements(p, child, base_dir)
                    elif child.name:
                        p.add_run(child.get_text())
                        
            elif element.name == 'table':
                # 表格
                rows = element.find_all('tr')
                if rows:
                    # 获取列数
                    first_row = rows[0]
                    col_count = len(first_row.find_all(['th', 'td']))
                    
                    if col_count > 0:
                        table = doc.add_table(rows=len(rows), cols=col_count)
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row in enumerate(rows):
                            cells = row.find_all(['th', 'td'])
                            for col_idx, cell in enumerate(cells):
                                if col_idx < col_count:
                                    table.rows[row_idx].cells[col_idx].text = cell.get_text()
                                    # 标题行加粗
                                    if row_idx == 0 and cell.name == 'th':
                                        for para in table.rows[row_idx].cells[col_idx].paragraphs:
                                            for run in para.runs:
                                                run.bold = True
                                        
            elif element.name == 'img':
                # 图片
                src = element.get('src', '')
                alt = element.get('alt', '')
                if src:
                    try:
                        # 处理相对路径
                        if base_dir and not src.startswith(('http://', 'https://', 'file://', 'data:')):
                            if src.startswith('file://'):
                                src = src[7:]
                            img_path = os.path.normpath(os.path.join(base_dir, src))
                        else:
                            img_path = src
                            
                        if os.path.exists(img_path):
                            doc.add_picture(img_path, width=Inches(5.0))
                            # 添加图片说明
                            if alt:
                                p = doc.add_paragraph(alt)
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.runs[0].italic = True
                                p.runs[0].font.size = Pt(9)
                    except Exception as e:
                        # 如果图片加载失败，添加占位文本
                        p = doc.add_paragraph(f'[图片: {alt or src}]')
                        p.italic = True
        
        doc.save(output_path)
        return True
    
    @staticmethod
    def _parse_inline_elements(paragraph, element, base_dir=None):
        """解析行内元素"""
        for child in element.children:
            if child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)
            elif child.name == 'a':
                run = paragraph.add_run(child.get_text())
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            elif child.name == 'br':
                paragraph.add_run('\n')
            elif child.name is None:
                # 纯文本
                if child.string:
                    paragraph.add_run(child.string)
            elif child.name:
                # 递归处理嵌套元素
                MarkdownConverter._parse_inline_elements(paragraph, child, base_dir)
    
    @staticmethod
    def md_to_html(md_text, output_path):
        """将 Markdown 转换为 HTML 文件"""
        # 转换为 HTML
        html_body = markdown.markdown(
            md_text,
            extensions=[
                'pymdownx.extra',
                'pymdownx.highlight',
                'pymdownx.superfences',
                'tables',
                'fenced_code',
                'toc',
                'meta'
            ]
        )
        
        # 构建完整的 HTML 文档
        html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>markFlet 导出文档</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Microsoft YaHei", sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h3 {{ font-size: 1.25em; }}
        code {{
            background-color: #f6f8fa;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            font-size: 85%;
        }}
        pre {{
            background-color: #f6f8fa;
            padding: 16px;
            overflow: auto;
            border-radius: 6px;
            line-height: 1.45;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        blockquote {{
            border-left: 4px solid #dfe2e5;
            padding-left: 16px;
            margin-left: 0;
            color: #6a737d;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 16px 0;
        }}
        th, td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        th {{
            background-color: #f6f8fa;
            font-weight: 600;
        }}
        tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        a {{
            color: #0366d6;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        ul, ol {{
            padding-left: 2em;
        }}
        li + li {{
            margin-top: 0.25em;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
        return True
    
    @staticmethod
    def md_to_pdf(md_text, output_path, base_dir=None):
        """将 Markdown 转换为 PDF 文档"""
        # 首先生成 HTML
        html_content = markdown.markdown(
            md_text,
            extensions=[
                'pymdownx.extra',
                'pymdownx.highlight',
                'pymdownx.superfences',
                'tables',
                'fenced_code',
                'toc'
            ]
        )
        
        # 构建完整 HTML 文档
        html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: "Microsoft YaHei", "SimHei", sans-serif;
            line-height: 1.6;
            padding: 20px;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 20px;
            margin-bottom: 10px;
        }}
        code {{
            background-color: #f5f5f5;
            padding: 2px 5px;
            border-radius: 3px;
            font-family: Consolas, monospace;
        }}
        pre {{
            background-color: #f5f5f5;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 3px solid #ccc;
            padding-left: 15px;
            margin-left: 0;
            color: #666;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
        }}
        th {{
            background-color: #f5f5f5;
        }}
        img {{
            max-width: 100%;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        # 使用 weasyprint 或 pdfkit 转换为 PDF
        if WEASYPRINT_AVAILABLE:
            WeasyHTML(string=html_template).write_pdf(output_path)
            return True
        elif PDFKIT_AVAILABLE:
            pdfkit.from_string(html_template, output_path)
            return True
        else:
            raise ImportError("未找到 PDF 库，请安装 weasyprint 或 pdfkit")


def main(page: ft.Page):
    """主应用"""
    page.title = "markFlet - Markdown 阅读器"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.padding = 0
    page.window_width = 1400
    page.window_height = 900
    
    # 设置全局字体为微软雅黑
    page.theme = ft.Theme(font_family="Microsoft YaHei")
    
    # 初始化数据库
    db = Database()
    
    # 当前打开的文件路径
    current_file = None
    
    # Markdown 转换器
    converter = MarkdownConverter()
    
    # 图片保存目录
    images_dir = None
    
    # 创建 UI 组件
    # 编辑区
    editor = ft.TextField(
        multiline=True,
        min_lines=30,
        max_lines=100,
        expand=True,
        border_color=ft.Colors.TRANSPARENT,
        bgcolor=ft.Colors.WHITE,
        text_size=14,
        text_style=ft.TextStyle(font_family="Microsoft YaHei"),
        on_change=lambda e: update_preview()
    )
    
    # 预览区 - 使用 Container 包裹确保正确展开
    preview_content = ft.Markdown(
        selectable=True,
        expand=True,
        extension_set=ft.MarkdownExtensionSet.GITHUB_WEB,
        on_tap_link=lambda e: page.launch_url(e.data)
    )
    
    # 将 Markdown 包裹在 Container 中，确保占满空间
    preview = ft.Container(
        content=preview_content,
        expand=True,
        padding=15,
    )
    
    # 状态栏
    status_text = ft.Text("就绪", size=12, font_family="Microsoft YaHei")
    
    # SnackBar（flet 0.80+ 用法）
    snack_bar = ft.SnackBar(content=ft.Text("", font_family="Microsoft YaHei"))
    page.overlay.append(snack_bar)
    
    def show_snack(message):
        """显示提示消息"""
        snack_bar.content.value = message
        snack_bar.open = True
        page.update()
    
    # 文件选择对话框（flet 0.80+ 中使用 page.services）
    file_picker = ft.FilePicker()
    page.services.append(file_picker)
    
    # 保存对话框（flet 0.80+ 中复用同一个 FilePicker 实例）
    save_file_dialog = file_picker
    
    # 预览区滚动引用
    preview_column_ref = ft.Ref[ft.Column]()
    
    import re
    
    def update_preview():
        """更新预览"""
        try:
            md_text = editor.value or ""
            
            # 处理图片路径：将相对路径转换为 file:// 绝对路径
            if current_file:
                base_dir = os.path.dirname(current_file)
                
                def replace_image_path(match):
                    alt_text = match.group(1)
                    img_path = match.group(2)
                    # 如果是相对路径且不是 http/https 链接
                    if not img_path.startswith(('http://', 'https://', 'file://')):
                        abs_path = os.path.normpath(os.path.join(base_dir, img_path))
                        # 使用 file:// 协议
                        if os.path.exists(abs_path):
                            return f'![{alt_text}](file://{abs_path})'
                    return match.group(0)
                
                # 替换 Markdown 图片语法中的路径
                md_text = re.sub(r'!\[(.*?)\]\((.*?)\)', replace_image_path, md_text)
            
            preview_content.value = md_text
            page.update()
        except Exception as e:
            status_text.value = f"预览更新失败: {str(e)}"
            page.update()
    
    def sync_preview_scroll(e: ft.OnScrollEvent):
        """同步预览区滚动"""
        try:
            if preview_column_ref.current and e.extent_total > 0:
                # 计算滚动比例 (0.0 - 1.0)
                scroll_ratio = e.extent_before / e.extent_total
                # 应用到预览区
                preview_column_ref.current.scroll_to(ratio=scroll_ratio, duration=0)
        except:
            pass
    
    async def open_file(e):
        """打开文件"""
        result = await file_picker.pick_files(
            dialog_title="选择 Markdown 文件",
            allowed_extensions=["md", "markdown", "txt"]
        )
        if result and len(result) > 0:
            file_path = result[0].path
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                editor.value = content
                nonlocal current_file, images_dir
                current_file = file_path
                # 重置图片目录
                images_dir = None
                update_preview()
                db.add_recent_file(file_path)
                page.title = f"markFlet - {os.path.basename(file_path)}"
                status_text.value = f"已打开: {file_path}"
                page.update()
            except Exception as ex:
                show_snack(f"打开文件失败: {str(ex)}")
    
    async def save_file(e):
        """保存文件"""
        nonlocal current_file
        
        if current_file:
            try:
                with open(current_file, 'w', encoding='utf-8') as f:
                    f.write(editor.value or "")
                status_text.value = f"已保存: {current_file}"
                show_snack("文件已保存")
            except Exception as ex:
                show_snack(f"保存失败: {str(ex)}")
        else:
            await save_as_file(e)
    
    async def save_as_file(e):
        """另存为"""
        file_path = await save_file_dialog.save_file(
            dialog_title="保存 Markdown 文件",
            file_name="untitled.md",
            allowed_extensions=["md"]
        )
        if file_path:
            if not file_path.endswith('.md'):
                file_path += '.md'
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(editor.value or "")
                nonlocal current_file
                current_file = file_path
                page.title = f"markFlet - {os.path.basename(file_path)}"
                status_text.value = f"已保存: {file_path}"
                db.add_recent_file(file_path)
                show_snack("文件已保存")
                page.update()
            except Exception as ex:
                show_snack(f"保存失败: {str(ex)}")
    
    async def export_word(e):
        """导出 Word"""
        file_path = await save_file_dialog.save_file(
            dialog_title="导出 Word 文档",
            file_name="document.docx",
            allowed_extensions=["docx"]
        )
        if file_path:
            if not file_path.endswith('.docx'):
                file_path += '.docx'
            try:
                base_dir = os.path.dirname(current_file) if current_file else None
                converter.md_to_docx(editor.value or "", file_path, base_dir)
                status_text.value = f"已导出: {file_path}"
                show_snack("Word 文档已导出")
                page.update()
            except Exception as ex:
                show_snack(f"导出失败: {str(ex)}")
    
    async def export_html(e):
        """导出 HTML"""
        file_path = await save_file_dialog.save_file(
            dialog_title="导出 HTML 文件",
            file_name="document.html",
            allowed_extensions=["html"]
        )
        if file_path:
            if not file_path.endswith('.html'):
                file_path += '.html'
            try:
                converter.md_to_html(editor.value or "", file_path)
                status_text.value = f"已导出: {file_path}"
                show_snack("HTML 文件已导出")
                page.update()
            except Exception as ex:
                show_snack(f"导出失败: {str(ex)}")
    
    async def export_pdf(e):
        """导出 PDF"""
        file_path = await save_file_dialog.save_file(
            dialog_title="导出 PDF 文档",
            file_name="document.pdf",
            allowed_extensions=["pdf"]
        )
        if file_path:
            if not file_path.endswith('.pdf'):
                file_path += '.pdf'
            try:
                base_dir = os.path.dirname(current_file) if current_file else None
                converter.md_to_pdf(editor.value or "", file_path, base_dir)
                status_text.value = f"已导出: {file_path}"
                show_snack("PDF 文档已导出")
                page.update()
            except ImportError:
                show_snack("请先安装 PDF 库: pip install weasyprint 或 pip install pdfkit")
            except Exception as ex:
                show_snack(f"导出失败: {str(ex)}")
    
    def toggle_theme(e):
        """切换主题"""
        if page.theme_mode == ft.ThemeMode.LIGHT:
            page.theme_mode = ft.ThemeMode.DARK
            editor.bgcolor = ft.Colors.GREY_900
            editor.text_style = ft.TextStyle(color=ft.Colors.WHITE, font_family="Microsoft YaHei")
        else:
            page.theme_mode = ft.ThemeMode.LIGHT
            editor.bgcolor = ft.Colors.WHITE
            editor.text_style = ft.TextStyle(color=ft.Colors.BLACK, font_family="Microsoft YaHei")
        page.update()
    
    def new_file(e):
        """新建文件"""
        nonlocal current_file
        editor.value = ""
        current_file = None
        images_dir = None
        page.title = "markFlet - 未命名"
        update_preview()
        status_text.value = "新建文件"
        page.update()
    
    def get_images_directory():
        """获取图片保存目录"""
        nonlocal images_dir
        if images_dir:
            return images_dir
        if current_file:
            # 如果有当前文件，在其所在目录创建 images 子目录
            base_dir = os.path.dirname(current_file)
            img_dir = os.path.join(base_dir, "images")
        else:
            # 否则使用临时目录
            img_dir = os.path.join(os.path.expanduser("~"), "markFlet_images")
        
        if not os.path.exists(img_dir):
            os.makedirs(img_dir)
        images_dir = img_dir
        return img_dir
    
    def paste_image_from_clipboard(e=None):
        """从剪贴板粘贴图片"""
        try:
            # 尝试从剪贴板获取图片
            img = ImageGrab.grabclipboard()
            
            if img is None:
                show_snack("剪贴板中没有图片")
                return False
            
            if not isinstance(img, Image.Image):
                show_snack("剪贴板内容不是图片")
                return False
            
            # 生成图片文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            img_filename = f"image_{timestamp}.png"
            img_dir = get_images_directory()
            img_path = os.path.join(img_dir, img_filename)
            
            # 保存图片
            img.save(img_path, "PNG")
            
            # 计算相对路径（相对于当前文件）
            if current_file:
                base_dir = os.path.dirname(current_file)
                rel_path = os.path.relpath(img_path, base_dir).replace("\\", "/")
            else:
                rel_path = img_path.replace("\\", "/")
            
            # 在编辑器当前光标位置插入 Markdown 图片语法
            cursor_pos = editor.selection.start if editor.selection else len(editor.value or "")
            current_text = editor.value or ""
            
            # 构建 Markdown 图片语法
            img_markdown = f"![图片]({rel_path})"
            
            # 在光标位置插入
            new_text = current_text[:cursor_pos] + img_markdown + current_text[cursor_pos:]
            editor.value = new_text
            
            # 更新预览
            update_preview()
            
            status_text.value = f"图片已保存: {img_filename}"
            show_snack(f"图片已粘贴并保存到: {rel_path}")
            page.update()
            return True
            
        except Exception as ex:
            show_snack(f"粘贴图片失败: {str(ex)}")
            return False
    
    def on_editor_paste(e):
        """处理编辑器粘贴事件"""
        # 延迟执行，让文本粘贴先完成
        def check_and_paste_image():
            try:
                img = ImageGrab.grabclipboard()
                if img is not None and isinstance(img, Image.Image):
                    # 如果剪贴板中有图片，处理图片粘贴
                    paste_image_from_clipboard()
            except:
                pass
        
        # 使用 page.run_thread 延迟检查剪贴板
        import threading
        threading.Timer(0.1, check_and_paste_image).start()
    
    # 工具栏
    toolbar = ft.Row(
        [
            ft.ElevatedButton("新建", icon=ft.Icons.ADD, on_click=new_file),
            ft.ElevatedButton("打开", icon=ft.Icons.FOLDER_OPEN, on_click=open_file),
            ft.ElevatedButton("保存", icon=ft.Icons.SAVE, on_click=save_file),
            ft.ElevatedButton("另存为", icon=ft.Icons.SAVE_AS, on_click=save_as_file),
            ft.VerticalDivider(width=10),
            ft.ElevatedButton("导出 Word", icon=ft.Icons.DESCRIPTION, on_click=export_word),
            ft.ElevatedButton("导出 HTML", icon=ft.Icons.CODE, on_click=export_html),
            ft.ElevatedButton("导出 PDF", icon=ft.Icons.PICTURE_AS_PDF, on_click=export_pdf),
            ft.VerticalDivider(width=10),
            ft.ElevatedButton("粘贴图片", icon=ft.Icons.IMAGE, on_click=paste_image_from_clipboard, tooltip="从剪贴板粘贴图片 (Ctrl+Shift+V)"),
            ft.VerticalDivider(width=10),
            ft.IconButton(
                icon=ft.Icons.DARK_MODE if page.theme_mode == ft.ThemeMode.LIGHT else ft.Icons.LIGHT_MODE,
                on_click=toggle_theme,
                tooltip="切换主题"
            ),
        ],
        alignment=ft.MainAxisAlignment.START,
        spacing=10
    )
    
    # 主布局
    page.add(
        ft.Column(
            [
                # 工具栏
                ft.Container(
                    content=toolbar,
                    padding=10,
                    bgcolor=ft.Colors.SURFACE_CONTAINER
                ),
                
                # 编辑区和预览区
                ft.Row(
                    [
                        # 编辑区
                        ft.Container(
                            content=ft.Column(
                                [
                                    ft.Text("编辑", weight=ft.FontWeight.BOLD, size=12, font_family="Microsoft YaHei"),
                                    editor
                                ],
                                expand=True,
                                scroll=ft.ScrollMode.AUTO,
                                on_scroll=lambda e: sync_preview_scroll(e)
                            ),
                            expand=True,
                            padding=10,
                            border=ft.border.all(1, ft.Colors.OUTLINE)
                        ),
                        
                        # 预览区
                        ft.Container(
                            content=ft.Column(
                                [
                                    ft.Text("预览", weight=ft.FontWeight.BOLD, size=12, font_family="Microsoft YaHei"),
                                    ft.Container(
                                        content=preview,
                                        expand=True,
                                        clip_behavior=ft.ClipBehavior.ANTI_ALIAS
                                    )
                                ],
                                expand=True,
                                spacing=5,
                                scroll=ft.ScrollMode.AUTO,
                                ref=preview_column_ref
                            ),
                            expand=True,
                            padding=10,
                            border=ft.border.all(1, ft.Colors.OUTLINE),
                            bgcolor=ft.Colors.WHITE if page.theme_mode == ft.ThemeMode.LIGHT else ft.Colors.GREY_900
                        )
                    ],
                    expand=True
                ),
                
                # 状态栏
                ft.Container(
                    content=status_text,
                    padding=10,
                    bgcolor=ft.Colors.SURFACE_CONTAINER
                )
            ],
            expand=True
        )
    )
    
    # 初始化预览
    editor.value = """# 欢迎使用 markFlet

这是一个基于 **Python Flet** 的 Markdown 阅读器。

## 功能特性

- 📝 Markdown 编辑
- 👁️ 实时预览
- 📄 Word 导出（支持图片、表格、代码块）
- 🌐 HTML 导出
- 📑 PDF 导出
- 🎨 主题切换
- 🖼️️ 图片粘贴支持

## 开始使用

1. 点击"打开"按钮加载 Markdown 文件
2. 在左侧编辑，右侧实时预览
3. 点击"导出 Word"生成 Word 文档

## 图片粘贴

点击工具栏的"粘贴图片"按钮，或复制图片后点击按钮，即可将图片插入到文档中。

## 代码示例

```python
print("Hello, markFlet!")
```

> 💡 **提示**：支持标准 Markdown 语法，包括标题、列表、代码块、引用、图片等。
"""
    update_preview()


if __name__ == "__main__":
    ft.run(main)