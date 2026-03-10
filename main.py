import flet as ft
import markdown
import sqlite3
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Database:
    """SQLite 数据库管理"""
    
    def __init__(self, db_path="markflet.db"):
        self.db_path = db_path
        self.init_db()
    
    def init_db(self):
        """初始化数据库"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS recent_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_path TEXT UNIQUE,
                file_name TEXT,
                opened_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()
        conn.close()
    
    def add_recent_file(self, file_path):
        """添加最近打开的文件"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        file_name = os.path.basename(file_path)
        cursor.execute('''
            INSERT OR REPLACE INTO recent_files (file_path, file_name, opened_at)
            VALUES (?, ?, CURRENT_TIMESTAMP)
        ''', (file_path, file_name))
        conn.commit()
        conn.close()
    
    def get_recent_files(self, limit=10):
        """获取最近打开的文件列表"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT file_path, file_name, opened_at 
            FROM recent_files 
            ORDER BY opened_at DESC 
            LIMIT ?
        ''', (limit,))
        files = cursor.fetchall()
        conn.close()
        return files


class MarkdownConverter:
    """Markdown 转换器"""
    
    @staticmethod
    def md_to_html(md_text):
        """将 Markdown 转换为 HTML"""
        return markdown.markdown(
            md_text,
            extensions=[
                'pymdownx.extra',
                'pymdownx.highlight',
                'pymdownx.superfences',
                'tables',
                'fenced_code',
            ]
        )
    
    @staticmethod
    def md_to_docx(md_text, output_path):
        """将 Markdown 转换为 Word 文档"""
        doc = Document()
        
        # 解析 Markdown
        lines = md_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # 处理标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                i += 1
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
                i += 1
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
                i += 1
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
                i += 1
            
            # 处理代码块
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # 跳过结束标记
                
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.5)
            
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                    items.append(lines[i].strip()[2:])
                    i += 1
                for item in items:
                    p = doc.add_paragraph(item, style='List Bullet')
            
            # 处理有序列表
            elif line[0:2].replace('.', '').isdigit():
                items = []
                while i < len(lines) and lines[i].strip() and lines[i].strip()[0:2].replace('.', '').isdigit():
                    item_text = lines[i].strip()
                    if '. ' in item_text:
                        items.append(item_text.split('. ', 1)[1])
                    i += 1
                for item in items:
                    p = doc.add_paragraph(item, style='List Number')
            
            # 处理引用
            elif line.startswith('> '):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith('> '):
                    quote_lines.append(lines[i].strip()[2:])
                    i += 1
                p = doc.add_paragraph(' '.join(quote_lines))
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                run = p.runs[0] if p.runs else p.add_run()
                run.italic = True
            
            # 处理普通段落（支持粗体、斜体、行内代码）
            else:
                p = doc.add_paragraph()
                
                # 处理粗体和斜体
                parts = line.split('**')
                for idx, part in enumerate(parts):
                    if idx % 2 == 1:  # 粗体部分
                        p.add_run(part).bold = True
                    else:
                        # 处理斜体
                        sub_parts = part.split('*')
                        for sub_idx, sub_part in enumerate(sub_parts):
                            if sub_idx % 2 == 1:  # 斜体部分
                                p.add_run(sub_part).italic = True
                            else:
                                # 处理行内代码
                                code_parts = sub_part.split('`')
                                for code_idx, code_part in enumerate(code_parts):
                                    if code_idx % 2 == 1:  # 代码部分
                                        run = p.add_run(code_part)
                                        run.font.name = 'Courier New'
                                        run.font.size = Pt(10)
                                    else:
                                        if code_part:
                                            p.add_run(code_part)
                i += 1
        
        doc.save(output_path)
        return True


def main(page: ft.Page):
    """主应用"""
    page.title = "markFlet - Markdown 阅读器"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.padding = 0
    page.window_width = 1400
    page.window_height = 900
    
    # 初始化数据库
    db = Database()
    
    # 当前打开的文件路径
    current_file = None
    
    # Markdown 转换器
    converter = MarkdownConverter()
    
    # 创建 UI 组件
    # 编辑区
    editor = ft.TextField(
        multiline=True,
        min_lines=30,
        max_lines=100,
        expand=True,
        border_color=ft.colors.TRANSPARENT,
        bgcolor=ft.colors.WHITE,
        text_size=14,
        font_family="Consolas",
        on_change=lambda e: update_preview()
    )
    
    # 预览区
    preview = ft.Markdown(
        selectable=True,
        expand=True,
        extension_set=ft.MarkdownExtensionSet.GITHUB_WEB,
        on_tap_link=lambda e: page.launch_url(e.data)
    )
    
    # 状态栏
    status_text = ft.Text("就绪", size=12)
    
    # 文件选择对话框
    file_picker = ft.FilePicker()
    page.overlay.append(file_picker)
    
    # 保存对话框
    save_file_dialog = ft.FilePicker()
    page.overlay.append(save_file_dialog)
    
    def update_preview():
        """更新预览"""
        try:
            md_text = editor.value or ""
            preview.value = md_text
            page.update()
        except Exception as e:
            status_text.value = f"预览更新失败: {str(e)}"
            page.update()
    
    def open_file(e):
        """打开文件"""
        def on_result(e: ft.FilePickerResultEvent):
            nonlocal current_file
            if e.files:
                file_path = e.files[0].path
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    editor.value = content
                    current_file = file_path
                    update_preview()
                    db.add_recent_file(file_path)
                    page.title = f"markFlet - {os.path.basename(file_path)}"
                    status_text.value = f"已打开: {file_path}"
                    page.update()
                except Exception as ex:
                    page.show_snack_bar(
                        ft.SnackBar(content=ft.Text(f"打开文件失败: {str(ex)}"))
                    )
        
        file_picker.on_result = on_result
        file_picker.pick_files(
            dialog_title="选择 Markdown 文件",
            allowed_extensions=["md", "markdown", "txt"]
        )
    
    def save_file(e):
        """保存文件"""
        nonlocal current_file
        
        if current_file:
            try:
                with open(current_file, 'w', encoding='utf-8') as f:
                    f.write(editor.value or "")
                status_text.value = f"已保存: {current_file}"
                page.show_snack_bar(
                    ft.SnackBar(content=ft.Text("文件已保存"))
                )
                page.update()
            except Exception as ex:
                page.show_snack_bar(
                    ft.SnackBar(content=ft.Text(f"保存失败: {str(ex)}"))
                )
        else:
            save_as_file(e)
    
    def save_as_file(e):
        """另存为"""
        def on_result(e: ft.FilePickerResultEvent):
            nonlocal current_file
            if e.path:
                file_path = e.path
                if not file_path.endswith('.md'):
                    file_path += '.md'
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(editor.value or "")
                    current_file = file_path
                    page.title = f"markFlet - {os.path.basename(file_path)}"
                    status_text.value = f"已保存: {file_path}"
                    db.add_recent_file(file_path)
                    page.show_snack_bar(
                        ft.SnackBar(content=ft.Text("文件已保存"))
                    )
                    page.update()
                except Exception as ex:
                    page.show_snack_bar(
                        ft.SnackBar(content=ft.Text(f"保存失败: {str(ex)}"))
                    )
        
        save_file_dialog.on_result = on_result
        save_file_dialog.save_file(
            dialog_title="保存 Markdown 文件",
            file_name="untitled.md",
            allowed_extensions=["md"]
        )
    
    def export_word(e):
        """导出 Word"""
        def on_result(e: ft.FilePickerResultEvent):
            if e.path:
                file_path = e.path
                if not file_path.endswith('.docx'):
                    file_path += '.docx'
                try:
                    converter.md_to_docx(editor.value or "", file_path)
                    status_text.value = f"已导出: {file_path}"
                    page.show_snack_bar(
                        ft.SnackBar(content=ft.Text("Word 文档已导出"))
                    )
                    page.update()
                except Exception as ex:
                    page.show_snack_bar(
                        ft.SnackBar(content=ft.Text(f"导出失败: {str(ex)}"))
                    )
        
        save_file_dialog.on_result = on_result
        save_file_dialog.save_file(
            dialog_title="导出 Word 文档",
            file_name="document.docx",
            allowed_extensions=["docx"]
        )
    
    def toggle_theme(e):
        """切换主题"""
        if page.theme_mode == ft.ThemeMode.LIGHT:
            page.theme_mode = ft.ThemeMode.DARK
            editor.bgcolor = ft.colors.GREY_900
            editor.text_style = ft.TextStyle(color=ft.colors.WHITE)
        else:
            page.theme_mode = ft.ThemeMode.LIGHT
            editor.bgcolor = ft.colors.WHITE
            editor.text_style = ft.TextStyle(color=ft.colors.BLACK)
        page.update()
    
    def new_file(e):
        """新建文件"""
        nonlocal current_file
        editor.value = ""
        current_file = None
        page.title = "markFlet - 未命名"
        update_preview()
        status_text.value = "新建文件"
        page.update()
    
    # 工具栏
    toolbar = ft.Row(
        [
            ft.ElevatedButton("新建", icon=ft.icons.ADD, on_click=new_file),
            ft.ElevatedButton("打开", icon=ft.icons.FOLDER_OPEN, on_click=open_file),
            ft.ElevatedButton("保存", icon=ft.icons.SAVE, on_click=save_file),
            ft.ElevatedButton("另存为", icon=ft.icons.SAVE_AS, on_click=save_as_file),
            ft.VerticalDivider(width=10),
            ft.ElevatedButton("导出 Word", icon=ft.icons.DESCRIPTION, on_click=export_word),
            ft.VerticalDivider(width=10),
            ft.IconButton(
                icon=ft.icons.DARK_MODE if page.theme_mode == ft.ThemeMode.LIGHT else ft.icons.LIGHT_MODE,
                on_click=toggle_theme,
                tooltip="切换主题"
            ),
        ],
        alignment=ft.MainAxisAlignment.START,
        spacing=10
    )
    
    # 主布局
    page.add(
        ft.Column(
            [
                # 工具栏
                ft.Container(
                    content=toolbar,
                    padding=10,
                    bgcolor=ft.colors.SURFACE_VARIANT
                ),
                
                # 编辑区和预览区
                ft.Row(
                    [
                        # 编辑区
                        ft.Container(
                            content=ft.Column(
                                [
                                    ft.Text("编辑", weight=ft.FontWeight.BOLD, size=12),
                                    editor
                                ],
                                expand=True
                            ),
                            expand=True,
                            padding=10,
                            border=ft.border.all(1, ft.colors.OUTLINE)
                        ),
                        
                        # 预览区
                        ft.Container(
                            content=ft.Column(
                                [
                                    ft.Text("预览", weight=ft.FontWeight.BOLD, size=12),
                                    ft.Container(
                                        content=preview,
                                        expand=True,
                                        padding=10,
                                        bgcolor=ft.colors.WHITE if page.theme_mode == ft.ThemeMode.LIGHT else ft.colors.GREY_900
                                    )
                                ],
                                expand=True
                            ),
                            expand=True,
                            padding=10,
                            border=ft.border.all(1, ft.colors.OUTLINE)
                        )
                    ],
                    expand=True
                ),
                
                # 状态栏
                ft.Container(
                    content=status_text,
                    padding=10,
                    bgcolor=ft.colors.SURFACE_VARIANT
                )
            ],
            expand=True
        )
    )
    
    # 初始化预览
    editor.value = """# 欢迎使用 markFlet

这是一个基于 **Python Flet** 的 Markdown 阅读器。

## 功能特性

- 📝 Markdown 编辑
- 👁️ 实时预览
- 📄 Word 导出
- 🎨 主题切换

## 开始使用

1. 点击"打开"按钮加载 Markdown 文件
2. 在左侧编辑，右侧实时预览
3. 点击"导出 Word"生成 Word 文档

## 代码示例

```python
print("Hello, markFlet!")
```

> 💡 **提示**：支持标准 Markdown 语法，包括标题、列表、代码块、引用等。
"""
    update_preview()


if __name__ == "__main__":
    ft.app(target=main)